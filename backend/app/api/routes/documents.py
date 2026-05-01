import os
import uuid
from typing import List
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update
from pydantic import BaseModel
from app.core.database import get_db
from app.core.dependencies import get_current_user
from app.core.config import (
    get_chunk_size,
    get_chunk_overlap,
    get_max_file_size,
    ALLOWED_FILE_TYPES,
    set_user_settings,
)
from app.core.embeddings import get_embeddings
from app.core.vector_store import insert_embeddings, delete_document_chunks
from app.models.models import User, Document, Chunk

router = APIRouter(prefix="/documents", tags=["文档"])

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

class DocumentResponse(BaseModel):
    id: str
    filename: str
    status: str
    created_at: str
    chunk_count: int

class PreviewResponse(BaseModel):
    filename: str
    content_type: str
    content: str
    file_url: str

def parse_text_from_file(file_path: str, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt" or ext == ".md":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    if ext == ".pdf":
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except ImportError:
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
                return "\n".join(text_parts)
            except ImportError:
                raise HTTPException(status_code=500, detail="缺少 PDF 解析依赖")

    if ext == ".docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return "\n".join([p.text for p in doc.paragraphs])

    if ext == ".csv":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    if ext == ".xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(file_path)
        texts = []
        for sheet in wb:
            for row in sheet.iter_rows(values_only=True):
                texts.append("\t".join(str(cell) for cell in row if cell is not None))
        return "\n".join(texts)

    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def split_text(text: str, chunk_size: int, overlap: int) -> List[str]:
    paragraphs = text.split("\n\n")
    chunks = []
    current_chunk = ""

    for para in paragraphs:
        if len(current_chunk) + len(para) > chunk_size and current_chunk:
            chunks.append(current_chunk.strip())
            current_chunk = para
        else:
            current_chunk += "\n\n" + para if current_chunk else para

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    if overlap > 0 and len(chunks) > 1:
        overlapped = []
        for i, chunk in enumerate(chunks):
            if i > 0:
                overlap_text = chunks[i - 1][-overlap:]
                chunk = overlap_text + "\n" + chunk
            overlapped.append(chunk)
        return overlapped

    return chunks

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ALLOWED_FILE_TYPES:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {ext}")

    file_content = await file.read()
    max_size = get_max_file_size(current_user.id)
    if len(file_content) > max_size:
        raise HTTPException(status_code=400, detail=f"文件大小超过 {max_size // (1024*1024)}MB 限制")

    file_id = str(uuid.uuid4())
    safe_filename = f"{file_id}{ext}"
    file_path = os.path.join(UPLOAD_DIR, safe_filename)

    with open(file_path, "wb") as f:
        f.write(file_content)

    doc = Document(
        user_id=current_user.id,
        filename=file.filename,
        file_url=f"/uploads/{safe_filename}",
        status="processing",
        file_size=len(file_content),
    )
    db.add(doc)
    await db.flush()
    await db.refresh(doc)

    try:
        text = parse_text_from_file(file_path, file.filename)
        chunk_size = get_chunk_size(current_user.id)
        overlap = get_chunk_overlap(current_user.id)
        chunks = split_text(text, chunk_size, overlap)

        chunk_records = []
        for i, chunk_text in enumerate(chunks):
            chunk_records.append({
                "doc_id": str(doc.id),
                "content": chunk_text,
                "chunk_index": i,
            })

        embeddings = await get_embeddings([r["content"] for r in chunk_records], current_user.id)
        for i, embedding in enumerate(embeddings):
            chunk_records[i]["embedding"] = embedding

        await insert_embeddings(db, chunk_records)

        await db.execute(
            update(Document)
            .where(Document.id == doc.id)
            .values(status="completed", chunk_count=len(chunks))
        )
        await db.commit()
        await db.refresh(doc)

    except Exception as e:
        await db.execute(
            update(Document)
            .where(Document.id == doc.id)
            .values(status="failed")
        )
        await db.commit()
        raise HTTPException(status_code=500, detail=f"文档处理失败: {str(e)}")

    return {
        "id": str(doc.id),
        "filename": doc.filename,
        "status": doc.status,
        "created_at": doc.created_at.isoformat(),
        "chunk_count": doc.chunk_count,
    }

@router.get("/", response_model=List[DocumentResponse])
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Document)
        .where(Document.user_id == current_user.id)
        .order_by(Document.created_at.desc())
    )
    docs = result.scalars().all()
    return [
        {
            "id": str(d.id),
            "filename": d.filename,
            "status": d.status,
            "created_at": d.created_at.isoformat(),
            "chunk_count": d.chunk_count,
        }
        for d in docs
    ]

@router.delete("/{doc_id}")
async def delete_document(
    doc_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    await delete_document_chunks(db, doc_id)
    await db.delete(doc)
    await db.commit()

    if doc.file_url:
        file_path = os.path.join(UPLOAD_DIR, os.path.basename(doc.file_url))
        if os.path.exists(file_path):
            os.remove(file_path)

    return {"message": "文档已删除"}

@router.get("/{doc_id}/preview", response_model=PreviewResponse)
async def preview_document(
    doc_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    file_path = os.path.join(UPLOAD_DIR, os.path.basename(doc.file_url))
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")

    ext = os.path.splitext(doc.filename)[1].lower()
    file_url = doc.file_url

    if ext == ".pdf":
        return {
            "filename": doc.filename,
            "content_type": "pdf",
            "content": "",
            "file_url": file_url,
        }

    if ext in (".txt", ".md", ".csv"):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            content_type = "markdown" if ext == ".md" else "text"
            return {
                "filename": doc.filename,
                "content_type": content_type,
                "content": content,
                "file_url": file_url,
            }
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="gbk") as f:
                content = f.read()
            return {
                "filename": doc.filename,
                "content_type": "text",
                "content": content,
                "file_url": file_url,
            }

    if ext in (".docx", ".doc"):
        content = parse_text_from_file(file_path, doc.filename)
        return {
            "filename": doc.filename,
            "content_type": "text",
            "content": content,
            "file_url": file_url,
        }

    if ext == ".xlsx":
        content = parse_text_from_file(file_path, doc.filename)
        return {
            "filename": doc.filename,
            "content_type": "text",
            "content": content,
            "file_url": file_url,
        }

    if ext == ".pptx":
        return {
            "filename": doc.filename,
            "content_type": "unsupported",
            "content": "暂不支持预览 PPTX 文件",
            "file_url": file_url,
        }

    raise HTTPException(status_code=400, detail=f"不支持预览该文件类型: {ext}")
